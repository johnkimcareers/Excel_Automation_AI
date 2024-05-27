import os

from docx import Document
import pandas as pd
import json
import openai

pd.set_option('display.max_columns', None)
pd.set_option('display.max_rows', None)
api_key = os.getenv('API_KEY')

class Assignment:
    def __init__(self, rubric_file_name):
        # use to specify response format of AI
        self.template_rubric = {}
        # Note, the two are link lockstep
        self.student_responses = {}
        self.student_graded_rubrics = {}
        # Setup Rubric_template variable
        self.setup(rubric_file_name)

    def load_json(self, file_path):
        try:
            with open(file_path, 'r') as json_file:
                data = json.load(json_file)
            return data
        except FileNotFoundError:
            print(f"Error: {file_path} not found.")
            return None
        except json.JSONDecodeError:
            print(f"Error: Could not decode JSON from {file_path}.")
            return None

    def read_excel(self, file_name):
        col = 0
        df = pd.read_excel(file_name)
        height = df.shape[0]
        width = df.shape[1]
        return col, df, height, width

    def setup(self, file_name):
        col, df, height, width = self.read_excel(file_name)

        for row in range(2, height):
            if isinstance(df.iloc[row, col], str):
                if df.iloc[row, col][0:4] == 'Data':
                    question = df.iloc[row,col]
                    if question not in self.template_rubric:
                        self.template_rubric[question] = ''
                        self.student_responses[question] = ''
                        self.student_graded_rubrics[question] = ''

    def parse_assignment(self, file_name):
        assignment = Document(file_name)
        full_text = []
        for paragraph in assignment.paragraphs:
            text = paragraph.text
            if text is None:
                continue
            else:
                full_text.append(text)

        final_text = '\n'.join(full_text)
        json_string = json.dumps(self.template_rubric)
        prompt = 'Please categorize every word of the following text into the provided JSON format, ' \
                 'without using commas between words: ' + json_string

        response = openai.ChatCompletion.create(
            model='gpt-4-1106-preview',
            response_format={'type': 'json_object'},
            messages=[
                {'role': 'system', 'content': prompt},
                {'role': 'user', 'content': final_text},
            ],
            api_key=api_key
        )

        extracted_response = response['choices'][0]['message']['content']
        print('------------------------------Response-------------------------------------')
        extracted_student_response = json.loads(extracted_response)

        for key, value in extracted_student_response.items():
            self.student_responses[key] += value + '\n ---'

        print(self.student_responses)

    def parse_marked_rubric(self, file_name):
        col, df, height, width = self.read_excel(file_name)
        grade_col = 5
        abc_col = 6

        for row in range(2, height):
            if isinstance(df.iloc[row, col], str):
                if df.iloc[row, col][0:4] == 'Data':
                    question = df.iloc[row,col]
                    self.student_graded_rubrics[question] += f'---{question}'
            elif isinstance(df.iloc[row, col], int):
                self.student_graded_rubrics[question] += f' {df.iloc[row, grade_col]}: {df.iloc[row, abc_col]} \n'

    def create_example_doc(self):
        doc = Document()

        for key in self.student_responses.keys():
            responses = list(filter(None,self.student_responses[key].split('---')))
            rubrics = list(filter(None,self.student_graded_rubrics[key].split('---')[1:]))
            print(responses)
            for index in range(len(rubrics)):
                doc.add_paragraph(responses[index])
                print('responses ', responses[index])
                doc.add_paragraph(rubrics[index])
                print('rubric ', rubrics[index])
                doc.add_paragraph()
        doc.save('responses.docx')
    def get_marked_rubric(self):
        return self.student_graded_rubrics

    def get_question_keys(self):
        return self.question_keys
